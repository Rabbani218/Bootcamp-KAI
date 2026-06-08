import os
import sys

def build_final():
    input_md = "Laporan_Teknis_NusaRail_Extended.md"
    output_docx = "Laporan_Teknis_NusaRail_Final.docx"

    print("Step 1: Mengonversi Markdown ke Word dengan pypandoc...")
    import pypandoc
    extra_args = [
        '--toc',
        '--number-sections',
        '-V', 'geometry:margin=1in',
    ]
    pypandoc.convert_file(input_md, 'docx', outputfile=output_docx, extra_args=extra_args)

    print("Step 2: Membersihkan format dan membuat Cover Page dengan python-docx...")
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document(output_docx)
    
    # --- ADD COVER PAGE AT THE BEGINNING ---
    # We will insert paragraphs before the very first paragraph (which is usually the TOC title or first header)
    first_p = doc.paragraphs[0]
    
    # Helper to insert centered bold text
    def insert_cover_text(text, size_pt, bold=False, space_after=0):
        p = first_p.insert_paragraph_before()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    # Spacing from top
    insert_cover_text("", 12, space_after=100)
    
    # Title
    insert_cover_text("LAPORAN TEKNIS", 16, bold=True, space_after=10)
    insert_cover_text("ARSITEKTUR & IMPLEMENTASI SOURCE CODE LENGKAP", 18, bold=True, space_after=24)
    
    # Subtitle
    insert_cover_text("NusaRail Vision System", 24, bold=True, space_after=10)
    insert_cover_text("Sistem Peringatan Dini Perlintasan Kereta Api Cerdas Terdistribusi", 14, bold=False, space_after=10)
    insert_cover_text("Integrasi YOLOv8, ByteTrack, dan Gemini 1.5 Pro", 14, bold=False, space_after=80)
    
    # Author details
    insert_cover_text("Disusun Oleh:", 12, bold=False, space_after=10)
    insert_cover_text("Muhammad Abdurrahman Rabbani", 14, bold=True, space_after=5)
    insert_cover_text("NIM: 15240969", 12, bold=False, space_after=40)
    
    insert_cover_text("Program Studi S1 Informatika", 14, bold=True, space_after=5)
    insert_cover_text("Fakultas Teknik & Informatika", 12, bold=False, space_after=5)
    insert_cover_text("Universitas Bina Sarana Informatika", 16, bold=True, space_after=10)
    insert_cover_text("Tahun 2026", 12, bold=False, space_after=20)
    
    # Page Break
    pb_p = first_p.insert_paragraph_before()
    run = pb_p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

    # --- REMOVE BLUE TEXT (FORCE BLACK) ---
    for style in doc.styles:
        if hasattr(style, 'font') and style.font.color:
            style.font.color.rgb = RGBColor(0, 0, 0)
            
    for para in doc.paragraphs:
        if para.style and para.style.font:
            para.style.font.color.rgb = RGBColor(0, 0, 0)
        for run in para.runs:
            if run.font and run.font.color and run.font.color.rgb:
                run.font.color.rgb = RGBColor(0, 0, 0)
            # Ensure font is a standard professional font
            if not run.font.name:
                run.font.name = 'Arial'

    doc.save(output_docx)
    print(f"Selesai! Laporan final yang sangat profesional tersimpan di: {output_docx}")

if __name__ == "__main__":
    build_final()
